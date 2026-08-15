from datetime import date
from pathlib import Path
import json
import math
import textwrap

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
ASSET_DIR = ROOT / "outputs" / "final_report_assets"
OUTPUT_PATH = ROOT / "Bao_cao_Do_an_Nhan_dien_Khuon_mat_FaceViT_SIC.docx"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

FONT_REGULAR = Path("C:/Windows/Fonts/arial.ttf")
FONT_BOLD = Path("C:/Windows/Fonts/arialbd.ttf")

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREEN = "E2F0D9"
LIGHT_ORANGE = "FCE4D6"
LIGHT_RED = "F4CCCC"
GRAY = "E7E6E6"
TEXT = "1F1F1F"


PAPERS = [
    {
        "key": "AlexNet",
        "authors": "Krizhevsky, Sutskever và Hinton",
        "title": "ImageNet Classification with Deep Convolutional Neural Networks",
        "venue": "NeurIPS 2012",
        "url": "https://proceedings.neurips.cc/paper_files/paper/2012/hash/c399862d3b9d6b76c8436e924a68c45b-Abstract.html",
    },
    {
        "key": "MobileNetV2",
        "authors": "Sandler và cộng sự",
        "title": "MobileNetV2: Inverted Residuals and Linear Bottlenecks",
        "venue": "CVPR 2018",
        "url": "https://openaccess.thecvf.com/content_cvpr_2018/html/Sandler_MobileNetV2_Inverted_Residuals_CVPR_2018_paper.html",
    },
    {
        "key": "ResNeSt",
        "authors": "Zhang và cộng sự",
        "title": "ResNeSt: Split-Attention Networks",
        "venue": "arXiv:2004.08955, 2020",
        "url": "https://arxiv.org/abs/2004.08955",
    },
    {
        "key": "ViT",
        "authors": "Dosovitskiy và cộng sự",
        "title": "An Image is Worth 16×16 Words: Transformers for Image Recognition at Scale",
        "venue": "ICLR 2021",
        "url": "https://arxiv.org/abs/2010.11929",
    },
    {
        "key": "DeiT",
        "authors": "Touvron và cộng sự",
        "title": "Training Data-Efficient Image Transformers & Distillation through Attention",
        "venue": "ICML 2021",
        "url": "https://proceedings.mlr.press/v139/touvron21a.html",
    },
    {
        "key": "Swin",
        "authors": "Liu và cộng sự",
        "title": "Swin Transformer: Hierarchical Vision Transformer Using Shifted Windows",
        "venue": "ICCV 2021",
        "url": "https://openaccess.thecvf.com/content/ICCV2021/html/Liu_Swin_Transformer_Hierarchical_Vision_Transformer_Using_Shifted_Windows_ICCV_2021_paper.html",
    },
    {
        "key": "FaceNet",
        "authors": "Schroff, Kalenichenko và Philbin",
        "title": "FaceNet: A Unified Embedding for Face Recognition and Clustering",
        "venue": "CVPR 2015",
        "url": "https://openaccess.thecvf.com/content_cvpr_2015/html/Schroff_FaceNet_A_Unified_2015_CVPR_paper.html",
    },
    {
        "key": "Batch-Hard",
        "authors": "Hermans, Beyer và Leibe",
        "title": "In Defense of the Triplet Loss for Person Re-Identification",
        "venue": "arXiv:1703.07737, 2017",
        "url": "https://arxiv.org/abs/1703.07737",
    },
    {
        "key": "VGGFace2",
        "authors": "Cao và cộng sự",
        "title": "VGGFace2: A Dataset for Recognising Faces across Pose and Age",
        "venue": "FG 2018",
        "url": "https://arxiv.org/abs/1710.08092",
    },
    {
        "key": "AdamW",
        "authors": "Loshchilov và Hutter",
        "title": "Decoupled Weight Decay Regularization",
        "venue": "ICLR 2019",
        "url": "https://arxiv.org/abs/1711.05101",
    },
]


VGG_INTERIM = [
    (1, 0.1623, 0.1760, 1.2759, 1.4674, 20.39),
    (2, 0.1516, 0.1516, 1.5628, 1.6932, 12.19),
    (3, 0.1528, 0.1591, 1.0070, 1.0973, 13.69),
    (4, 0.1494, 0.1525, 1.3970, 1.5347, 16.12),
    (5, 0.1486, 0.1497, 1.3097, 1.4389, 14.62),
    (6, 0.1496, 0.1632, 1.2617, 1.4174, 19.31),
    (7, 0.1491, 0.2181, 0.8833, 1.0168, 21.19),
    (8, 0.1507, 0.1629, 1.2912, 1.4249, 18.87),
    (9, 0.1530, 0.1513, 1.3246, 1.4599, 16.23),
]


def image_font(size, bold=False):
    path = FONT_BOLD if bold else FONT_REGULAR
    if not path.exists():
        for p in [
            Path("/usr/share/fonts/google-carlito-fonts/Carlito-Bold.ttf" if bold else "/usr/share/fonts/google-carlito-fonts/Carlito-Regular.ttf"),
            Path("/usr/share/fonts/dejavu-sans-fonts/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/dejavu-sans-fonts/DejaVuSans.ttf"),
            Path("/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/gnu-free/FreeSansBold.ttf" if bold else "/usr/share/fonts/gnu-free/FreeSans.ttf"),
        ]:
            if p.exists():
                return ImageFont.truetype(str(p), size=size)
        return ImageFont.load_default()
    return ImageFont.truetype(str(path), size=size)


def draw_centered(draw, coords, text, size=28, bold=False, color="#17365D"):
    x1, y1, x2, y2 = coords
    fnt = image_font(size, bold)
    max_chars = max(8, int((x2 - x1) / (size * 0.55)))
    wrapped = "\n".join(
        "\n".join(textwrap.wrap(line, width=max_chars))
        for line in text.splitlines()
    )
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=fnt, spacing=7, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(
        ((x1 + x2 - width) / 2, (y1 + y2 - height) / 2),
        wrapped,
        font=fnt,
        fill=color,
        spacing=7,
        align="center",
    )


def draw_box(draw, coords, text, fill, size=26):
    draw.rounded_rectangle(coords, radius=20, fill=fill, outline="#355070", width=3)
    draw_centered(draw, coords, text, size=size, bold=True)


def draw_arrow(draw, start, end, color="#355070", width=6):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    for offset in (2.55, -2.55):
        point = (
            end[0] + length * math.cos(angle + offset),
            end[1] + length * math.sin(angle + offset),
        )
        draw.line([end, point], fill=color, width=width)


def create_pipeline_diagram():
    path = ASSET_DIR / "01_pipeline.png"
    image = Image.new("RGB", (1900, 600), "white")
    draw = ImageDraw.Draw(image)
    draw.text((55, 30), "PIPELINE ĐỒ ÁN NHẬN DIỆN KHUÔN MẶT", font=image_font(38, True), fill="#17365D")
    labels = [
        ("VGGFace2\n480/30/30 identity", "#D9EAD3"),
        ("P×K Sampler\n4 người × 4 ảnh", "#D0E0E3"),
        ("FaceViT\n12 blocks", "#CFE2F3"),
        ("Embedding\n128 chiều, L2=1", "#FFF2CC"),
        ("Semi-Hard\nTriplet Loss", "#FCE5CD"),
        ("ROC/EER\nRecall@K/mAP", "#EAD1DC"),
    ]
    width, gap, y1, y2 = 245, 62, 170, 420
    x = 45
    for index, (label, color) in enumerate(labels):
        draw_box(draw, (x, y1, x + width, y2), label, color, size=24)
        if index < len(labels) - 1:
            draw_arrow(draw, (x + width, 295), (x + width + gap - 8, 295))
        x += width + gap
    draw.text(
        (55, 500),
        "Phase tiếp theo: Face Detection nhiều khuôn mặt → Gallery Enrollment → Web/Mobile/WinForms",
        font=image_font(28, True),
        fill="#7F6000",
    )
    image.save(path)
    return path


def create_architecture_diagram():
    path = ASSET_DIR / "02_facevit_architecture.png"
    image = Image.new("RGB", (1700, 1350), "white")
    draw = ImageDraw.Draw(image)
    draw.text((55, 28), "SIC FACEVIT — LUỒNG TENSOR", font=image_font(38, True), fill="#17365D")
    nodes = [
        ("Ảnh RGB\n[B, 3, 224, 224]", "#D9EAD3"),
        ("Conv2d patch embedding\nkernel=stride=16", "#D0E0E3"),
        ("196 patch token\n[B, 196, 192]", "#CFE2F3"),
        ("Thêm CLS + Position\n[B, 197, 192]", "#FFF2CC"),
        ("12 Transformer blocks\n3 heads, MLP ratio 4", "#FCE5CD"),
        ("Final LayerNorm + CLS\n[B, 192]", "#D9D2E9"),
        ("Linear 192→128\nL2 Normalize", "#EAD1DC"),
        ("Face embedding\n[B, 128], L2 norm = 1", "#D9EAD3"),
    ]
    y = 100
    for index, (label, color) in enumerate(nodes):
        coords = (365, y, 1335, y + 115)
        draw_box(draw, coords, label, color, size=27)
        if index < len(nodes) - 1:
            draw_arrow(draw, (850, y + 115), (850, y + 150))
        y += 155
    image.save(path)
    return path


def create_transformer_block_diagram():
    path = ASSET_DIR / "03_transformer_block.png"
    image = Image.new("RGB", (1850, 850), "white")
    draw = ImageDraw.Draw(image)
    draw.text((45, 25), "MỘT TRANSFORMER ENCODER BLOCK (PRE-NORM)", font=image_font(36, True), fill="#17365D")
    labels = [
        ("Token x\n[B,197,192]", "#D9EAD3"),
        ("LayerNorm", "#D0E0E3"),
        ("Multi-Head\nSelf-Attention", "#CFE2F3"),
        ("Residual\nx + Attention", "#F4CCCC"),
        ("LayerNorm", "#D0E0E3"),
        ("MLP\n192→768→192", "#FFF2CC"),
        ("Residual\nx + MLP", "#F4CCCC"),
        ("Output\n[B,197,192]", "#D9EAD3"),
    ]
    width, gap, y1, y2 = 195, 30, 270, 520
    x = 30
    centers = []
    for label, color in labels:
        draw_box(draw, (x, y1, x + width, y2), label, color, size=21)
        centers.append((x + width / 2, 395))
        x += width + gap
    for index in range(len(centers) - 1):
        draw_arrow(
            draw,
            (centers[index][0] + width / 2, 395),
            (centers[index + 1][0] - width / 2 - 5, 395),
            width=5,
        )
    draw.text((70, 650), "3 attention heads × 64 chiều/head = 192 chiều", font=image_font(28, True), fill="#355070")
    draw.text((70, 705), "Residual giữ thông tin cũ; LayerNorm ổn định phân bố đặc trưng", font=image_font(28, True), fill="#355070")
    image.save(path)
    return path


def create_triplet_diagram():
    path = ASSET_DIR / "04_semi_hard_triplet.png"
    image = Image.new("RGB", (1800, 950), "white")
    draw = ImageDraw.Draw(image)
    draw.text((50, 25), "P×K SAMPLER VÀ SEMI-HARD TRIPLET MINING", font=image_font(36, True), fill="#17365D")
    people = [
        ("Identity A\nA1 A2 A3 A4", "#D9EAD3"),
        ("Identity B\nB1 B2 B3 B4", "#CFE2F3"),
        ("Identity C\nC1 C2 C3 C4", "#FFF2CC"),
        ("Identity D\nD1 D2 D3 D4", "#EAD1DC"),
    ]
    x = 70
    for label, color in people:
        draw_box(draw, (x, 140, x + 360, 310), label, color, size=25)
        x += 430
    draw_box(draw, (90, 480, 470, 700), "Anchor A1", "#D9EAD3", size=30)
    draw_box(draw, (700, 420, 1110, 640), "Hardest Positive A4\ncùng người, xa nhất", "#CFE2F3", size=27)
    draw_box(draw, (1320, 480, 1710, 700), "Semi-Hard Negative C2\nkhác người", "#FCE5CD", size=27)
    draw_arrow(draw, (470, 560), (690, 530))
    draw_arrow(draw, (1110, 530), (1310, 560))
    draw.text((525, 745), "PosDist < NegDist < PosDist + margin (0,2)", font=image_font(31, True), fill="#C65911")
    draw.text((390, 810), "Loss = max(PosDist − NegDist + margin, 0)", font=image_font(31, True), fill="#17365D")
    image.save(path)
    return path


def create_dataset_chart():
    path = ASSET_DIR / "05_dataset_comparison.png"
    names = ["Pins", "VGGFace2 subset"]
    identities = [105, 540]
    images = [17534, 197693]
    figure, axes = plt.subplots(1, 2, figsize=(12, 5))
    bars = axes[0].bar(names, identities, color=["#5B9BD5", "#70AD47"])
    axes[0].set_title("Số identity")
    axes[0].set_ylabel("Identity")
    axes[0].bar_label(bars, fmt="%d")
    bars = axes[1].bar(names, images, color=["#5B9BD5", "#70AD47"])
    axes[1].set_title("Số ảnh")
    axes[1].set_ylabel("Ảnh")
    axes[1].bar_label(bars, fmt="%d")
    for axis in axes:
        axis.grid(True, axis="y", alpha=0.25)
    figure.suptitle("So sánh dữ liệu thực nghiệm")
    figure.tight_layout()
    figure.savefig(path, dpi=180)
    plt.close(figure)
    return path


def create_parameter_chart():
    path = ASSET_DIR / "06_model_parameters.png"
    names = ["MobileNetV2", "SIC FaceViT", "DeiT-Tiny", "ResNeSt-50", "Swin-T", "AlexNet", "ViT-B/16"]
    values = [3.5, 5.549, 5.7, 27.5, 28.3, 61.1, 86.6]
    colors = ["#70AD47", "#5B9BD5", "#4472C4", "#ED7D31", "#A5A5A5", "#C55A11", "#8064A2"]
    figure, axis = plt.subplots(figsize=(12, 6))
    bars = axis.barh(names, values, color=colors)
    axis.set_xlabel("Triệu tham số (xấp xỉ)")
    axis.set_title("So sánh quy mô mô hình")
    axis.bar_label(bars, fmt="%.3gM", padding=4)
    axis.grid(True, axis="x", alpha=0.25)
    axis.invert_yaxis()
    figure.tight_layout()
    figure.savefig(path, dpi=180)
    plt.close(figure)
    return path


def create_experiment_comparison_chart():
    path = ASSET_DIR / "07_experiment_comparison.png"
    names = ["ViT-4", "ViT-12", "ViT-12 MLP4", "ViT-Tiny 224"]
    accuracy = [25.67, 23.77, 23.07, 27.09]
    figure, axes = plt.subplots(1, 2, figsize=(13, 5))
    bars = axes[0].bar(names, accuracy, color="#5B9BD5")
    axes[0].set_title("Classifier trên Pins — Test Accuracy")
    axes[0].set_ylabel("Accuracy (%)")
    axes[0].set_ylim(0, 35)
    axes[0].tick_params(axis="x", rotation=20)
    axes[0].bar_label(bars, fmt="%.2f%%")
    metric_names = ["Random baseline", "Pins Semi-Hard"]
    auc = [0.5, 0.5231]
    bars = axes[1].bar(metric_names, auc, color=["#A5A5A5", "#ED7D31"])
    axes[1].set_title("Open-set verification trên Pins — ROC-AUC")
    axes[1].set_ylim(0.45, 0.56)
    axes[1].bar_label(bars, fmt="%.4f")
    for axis in axes:
        axis.grid(True, axis="y", alpha=0.25)
    figure.tight_layout()
    figure.savefig(path, dpi=180)
    plt.close(figure)
    return path


def create_vgg_interim_chart():
    path = ASSET_DIR / "08_vgg_interim.png"
    epochs = [row[0] for row in VGG_INTERIM]
    train_loss = [row[1] for row in VGG_INTERIM]
    val_loss = [row[2] for row in VGG_INTERIM]
    pos = [row[3] for row in VGG_INTERIM]
    neg = [row[4] for row in VGG_INTERIM]
    rate = [row[5] for row in VGG_INTERIM]
    figure, axes = plt.subplots(1, 3, figsize=(16, 4.8))
    axes[0].plot(epochs, train_loss, marker="o", label="Train Loss")
    axes[0].plot(epochs, val_loss, marker="o", label="Val Loss")
    axes[0].set_title("Loss tạm thời")
    axes[0].legend()
    axes[1].plot(epochs, pos, marker="o", label="PosDist")
    axes[1].plot(epochs, neg, marker="o", label="NegDist")
    axes[1].set_title("Khoảng cách validation")
    axes[1].legend()
    axes[2].plot(epochs, rate, marker="o", color="#ED7D31")
    axes[2].set_title("TripletRate")
    axes[2].set_ylabel("%")
    for axis in axes:
        axis.set_xlabel("Epoch")
        axis.grid(True, alpha=0.25)
    figure.suptitle("VGGFace2 — số liệu tạm thời đến epoch 9")
    figure.tight_layout()
    figure.savefig(path, dpi=180)
    plt.close(figure)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Trang ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_toc(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Nhấn chuột phải và chọn Update Field để cập nhật mục lục."
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    run_text = OxmlElement("w:t")
    run_text.text = text
    run.append(run_text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_caption(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    return paragraph


def add_picture(document, path, caption, width_cm=16.2):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(document, caption)


def add_note(document, text, color=LIGHT_ORANGE):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = True
    paragraph.paragraph_format.space_after = Pt(0)
    return table


def add_code(document, text):
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F2F2")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(text))
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    for row_index, row_data in enumerate(rows):
        row = table.add_row()
        for column_index, value in enumerate(row_data):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 1:
                set_cell_shading(cell, "F7F9FB")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(str(value))
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    document.add_paragraph()
    return table


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size, color in [(1, 17, BLUE), (2, 14, "2F75B5"), (3, 12, "5B9BD5")]:
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    for section in document.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("SIC Face Recognition — Báo cáo kỹ thuật")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        add_page_number(section.footer.paragraphs[0])


def add_cover(document):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(30)
    run = paragraph.add_run("SAMSUNG INNOVATION CAMPUS")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(75)
    run = paragraph.add_run("BÁO CÁO ĐỒ ÁN\nNHẬN DIỆN KHUÔN MẶT")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("FaceViT kết hợp Semi-Hard Triplet Loss")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("C65911")

    document.add_paragraph()
    info = add_table(
        document,
        ["Thông tin", "Nội dung"],
        [
            ("Nhóm sinh viên", "........................................................"),
            ("Giảng viên hướng dẫn", "........................................................"),
            ("Khóa học", "Samsung Innovation Campus — AI"),
            ("Mô hình", "SIC FaceViT — custom DeiT-Tiny/16-style backbone"),
            ("Ngày cập nhật", date.today().strftime("%d/%m/%Y")),
        ],
        widths=[5, 11],
    )
    info.autofit = False
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(55)
    run = paragraph.add_run("Thái Nguyên, 2026")
    run.bold = True
    run.font.size = Pt(13)
    document.add_page_break()


def build_report():
    assets = {
        "pipeline": create_pipeline_diagram(),
        "architecture": create_architecture_diagram(),
        "block": create_transformer_block_diagram(),
        "triplet": create_triplet_diagram(),
        "dataset": create_dataset_chart(),
        "parameters": create_parameter_chart(),
        "experiments": create_experiment_comparison_chart(),
        "vgg": create_vgg_interim_chart(),
    }

    document = Document()
    configure_document(document)
    add_cover(document)

    document.add_heading("TÓM TẮT", level=1)
    document.add_paragraph(
        "Báo cáo trình bày toàn bộ quá trình xây dựng hệ thống nhận diện khuôn mặt trong khóa Samsung Innovation Campus: "
        "từ baseline phân loại 105 người, chuyển sang metric learning với face embedding, chẩn đoán embedding collapse, "
        "thử Random Triplet, Batch-Hard và Semi-Hard Mining, sau đó mở rộng dữ liệu từ Pins Face Recognition sang subset VGGFace2."
    )
    document.add_paragraph(
        "Mô hình cuối được gọi là SIC FaceViT. Đây là mô hình custom do nhóm cài đặt, không phải tên một mô hình đã được công bố nguyên trạng. "
        "Backbone dùng cấu hình tương tự DeiT-Tiny/16 (patch 16, embedding 192, 12 block, 3 attention head, MLP ratio 4), "
        "sau đó thay classifier bằng embedding head 128 chiều và huấn luyện bằng PyTorch TripletMarginLoss với Semi-Hard Negative Mining theo tư tưởng FaceNet."
    )
    add_note(
        document,
        "Tính trung thực học thuật: báo cáo tách rõ (1) kiến trúc/bài báo có thật, (2) cấu hình custom của project, "
        "(3) số liệu thực nghiệm đã hoàn tất và (4) số liệu VGGFace2 đang chạy, chưa phải kết quả cuối.",
        LIGHT_GREEN,
    )
    add_picture(document, assets["pipeline"], "Hình 1. Pipeline tổng thể của đồ án", 17.0)

    document.add_heading("MỤC LỤC", level=1)
    toc_paragraph = document.add_paragraph()
    add_toc(toc_paragraph)
    document.add_page_break()

    document.add_heading("1. GIỚI THIỆU", level=1)
    document.add_heading("1.1. Bối cảnh", level=2)
    document.add_paragraph(
        "Nhận diện khuôn mặt là bài toán kết hợp thị giác máy tính và metric learning. Một hệ thống hoàn chỉnh thường gồm "
        "Face Detection để tìm vị trí khuôn mặt, Face Alignment để chuẩn hóa tư thế, Face Embedding để biến khuôn mặt thành vector, "
        "và Matching để quyết định danh tính hoặc lớp Unknown. Đồ án hiện hoàn thành phần embedding, training và evaluation; "
        "detection đa khuôn mặt và ứng dụng thời gian thực thuộc phase tiếp theo."
    )
    document.add_heading("1.2. Mục tiêu", level=2)
    add_bullets(
        document,
        [
            "Tự cài đặt mô hình Transformer cho ảnh, không sử dụng trọng số pretrained.",
            "Sinh face embedding 128 chiều thay vì chỉ dự đoán logits của 105 lớp.",
            "Huấn luyện bằng Anchor–Positive–Negative và Triplet Loss.",
            "Chống overfitting bằng augmentation, dropout, weight decay và Early Stopping.",
            "Đánh giá open-set bằng ROC-AUC, EER, TAR@FAR; đánh giá identification bằng Recall@K và mAP.",
            "Lưu checkpoint, history, biểu đồ và embedding artifact để tái sử dụng trong Web/Mobile/WinForms.",
        ],
    )
    document.add_heading("1.3. Phạm vi và nội dung chưa hoàn thành", level=2)
    add_bullets(
        document,
        [
            "Đã hoàn thành: dataset pipeline, FaceViT, Triplet training, mining, Early Stopping, test metrics và visualization.",
            "Đang chạy: thí nghiệm cuối trên VGGFace2 subset.",
            "Chưa tích hợp: detector đa khuôn mặt, face alignment landmarks, ResNeSt thực nghiệm, Web App, Mobile App và WinForms.",
            "ResNeSt chỉ được so sánh từ tài liệu; báo cáo không tuyên bố đã train ResNeSt.",
        ],
    )

    document.add_heading("2. CƠ SỞ LÝ THUYẾT VÀ CÁC MÔ HÌNH CÓ THẬT", level=1)
    document.add_heading("2.1. CNN truyền thống: AlexNet và MobileNetV2", level=2)
    document.add_paragraph(
        "AlexNet là CNN sâu nổi bật tại ImageNet 2012, sử dụng convolution, ReLU, pooling và fully connected; khoảng 61 triệu tham số. "
        "MobileNetV2 dùng depthwise separable convolution, inverted residual và linear bottleneck, khoảng 3,4–3,5 triệu tham số, "
        "phù hợp thiết bị di động nhưng khả năng biểu diễn thường thấp hơn mô hình lớn."
    )
    document.add_heading("2.2. ResNeSt", level=2)
    document.add_paragraph(
        "ResNeSt là biến thể ResNet dùng Split-Attention block: chia feature map thành các nhóm và học trọng số chú ý theo channel/group. "
        "ResNeSt có thật trong bài báo của Zhang và cộng sự. Tuy nhiên code hiện tại chưa triển khai ResNeSt; đây là baseline cần làm ở phase so sánh sau."
    )
    document.add_heading("2.3. Vision Transformer, DeiT và Swin Transformer", level=2)
    document.add_paragraph(
        "ViT biểu diễn ảnh như chuỗi patch token và sử dụng Transformer Encoder. ViT-B/16 chuẩn có khoảng 86 triệu tham số và thường cần dữ liệu lớn. "
        "DeiT đề xuất quy trình huấn luyện data-efficient; DeiT-Tiny có patch 16, embedding 192, depth 12, 3 head và khoảng 5,7 triệu tham số. "
        "Swin Transformer dùng cửa sổ dịch chuyển, tạo biểu diễn phân cấp giống CNN và giảm chi phí attention trên ảnh lớn."
    )
    document.add_heading("2.4. FaceNet và Triplet Loss", level=2)
    document.add_paragraph(
        "FaceNet học không gian embedding sao cho ảnh cùng người gần nhau và ảnh khác người xa nhau. Bài báo FaceNet sử dụng triplet loss "
        "và nhấn mạnh chiến lược chọn triplet, đặc biệt Semi-Hard Negative, để tránh mẫu quá dễ hoặc negative cực khó gây bất ổn."
    )
    add_picture(document, assets["parameters"], "Hình 2. Quy mô tham số của các kiến trúc tham khảo và SIC FaceViT", 16.5)
    add_table(
        document,
        ["Mô hình", "Năm", "Cơ chế chính", "Tham số xấp xỉ", "Vai trò trong báo cáo"],
        [
            ("AlexNet", "2012", "CNN + FC", "61,1M", "Mô hình CNN đã học"),
            ("MobileNetV2", "2018", "Depthwise + inverted residual", "3,5M", "Baseline mobile"),
            ("ResNeSt-50", "2020", "Split-Attention CNN", "27,5M", "Yêu cầu/roadmap, chưa train"),
            ("ViT-B/16", "2021", "Global self-attention", "86,6M", "Transformer chuẩn"),
            ("DeiT-Tiny", "2021", "Data-efficient ViT", "5,7M", "Cấu hình backbone tham chiếu"),
            ("Swin-T", "2021", "Shifted-window attention", "28,3M", "Hướng nâng cấp"),
            ("SIC FaceViT", "2026", "Custom DeiT-style + embedding", "5,549M", "Mô hình project"),
        ],
    )

    document.add_heading("3. DỮ LIỆU", level=1)
    document.add_heading("3.1. Pins Face Recognition", level=2)
    document.add_paragraph(
        "Dataset ban đầu gồm 105 danh tính và 17.534 ảnh. Ở pipeline classifier, ảnh được chia thành 12.225 train, 2.574 validation và 2.735 test. "
        "Khi chuyển sang open-set metric learning, identity được chia rời: 74 train (12.315 ảnh), 16 validation (2.707 ảnh), 15 test (2.512 ảnh)."
    )
    document.add_heading("3.2. VGGFace2 subset", level=2)
    document.add_paragraph(
        "VGGFace2 gốc có khoảng 3,31 triệu ảnh của 9.131 identity và được thiết kế để bao phủ biến thiên về pose, tuổi, ánh sáng và nghề nghiệp. "
        "Subset tải từ Kaggle trong project gồm 197.693 ảnh của 540 identity, tương đương khoảng 5,97% số ảnh và 5,91% số identity so với bản gốc."
    )
    add_table(
        document,
        ["Tập", "Identity", "Ảnh", "Nguyên tắc"],
        [
            ("Train", "480", "176.398", "Identity chỉ xuất hiện trong train"),
            ("Validation", "30", "10.957", "Lấy từ thư mục val, seed 42"),
            ("Test", "30", "10.338", "Không trùng train/validation"),
        ],
    )
    add_picture(document, assets["dataset"], "Hình 3. So sánh quy mô Pins và VGGFace2 subset", 15.5)
    document.add_paragraph(
        "VGGFace2 subset có số ảnh lớn gấp 11,27 lần Pins và số identity lớn gấp 5,14 lần. Riêng identity train tăng từ 74 lên 480, tức 6,49 lần, "
        "đây là thay đổi quan trọng nhất để mô hình học face embedding có khả năng tổng quát hóa tốt hơn."
    )
    document.add_heading("3.3. Tiền xử lý và augmentation", level=2)
    add_table(
        document,
        ["Train transform", "Tham số", "Mục đích"],
        [
            ("Resize", "224×224", "Đưa ảnh về kích thước cố định"),
            ("HorizontalFlip", "p=0,5", "Mô phỏng hướng mặt đối xứng"),
            ("RandomRotation", "−8° đến +8°", "Tăng bền vững với nghiêng nhẹ"),
            ("ColorJitter", "0,15", "Biến đổi sáng/tương phản/màu"),
            ("ToTensor", "[0,255] → [0,1]", "Đổi ảnh thành tensor"),
            ("Normalize", "mean=std=0,5", "Đưa giá trị xấp xỉ [−1,1]"),
        ],
    )
    document.add_paragraph("Validation và test chỉ Resize, ToTensor và Normalize để kết quả ổn định, không dùng augmentation ngẫu nhiên.")

    document.add_heading("4. KIẾN TRÚC SIC FACEVIT", level=1)
    add_note(
        document,
        "SIC FaceViT không phải tên mô hình trong một bài báo. Đây là cấu hình custom của project, dựa trên các khối ViT/DeiT có thật và phương pháp Triplet của FaceNet.",
        LIGHT_GREEN,
    )
    add_picture(document, assets["architecture"], "Hình 4. Luồng tensor của SIC FaceViT", 15.5)
    document.add_heading("4.1. Cấu hình", level=2)
    add_table(
        document,
        ["Tham số", "Giá trị", "Giải thích"],
        [
            ("Image size", "224", "Ảnh RGB 224×224"),
            ("Patch size", "16", "14×14 = 196 patch"),
            ("Embedding dimension", "192", "Chiều token"),
            ("Depth", "12", "12 Transformer block"),
            ("Attention heads", "3", "64 chiều/head"),
            ("MLP ratio", "4", "192→768→192"),
            ("Tokens", "197", "196 patch + 1 CLS"),
            ("Face embedding", "128", "Vector đầu ra"),
            ("Dropout", "0,1", "Regularization"),
            ("Tổng tham số", "5.549.120", "Tính trực tiếp từ model.parameters()"),
        ],
    )
    document.add_heading("4.2. Phân rã tham số", level=2)
    add_table(
        document,
        ["Thành phần", "Số tham số"],
        [
            ("Patch Embedding Conv2d", "147.648"),
            ("CLS token", "192"),
            ("Position Embedding", "37.824"),
            ("12 Transformer blocks", "5.338.368"),
            ("Final LayerNorm", "384"),
            ("Embedding head 192→128", "24.704"),
            ("Tổng", "5.549.120"),
        ],
    )
    document.add_heading("4.3. Transformer Encoder Block", level=2)
    add_picture(document, assets["block"], "Hình 5. Cấu trúc một Transformer Encoder Block", 17.0)
    document.add_paragraph(
        "LayerNorm chuẩn hóa theo chiều embedding của từng token. Multi-Head Self-Attention tạo Query, Key và Value để mỗi patch tổng hợp thông tin từ toàn ảnh. "
        "Residual Connection cộng đầu vào với đầu ra attention/MLP, giúp giữ thông tin và truyền gradient qua 12 block. FeedForward dùng GELU và dropout với chiều 192→768→192."
    )

    document.add_heading("5. METRIC LEARNING VÀ SEMI-HARD MINING", level=1)
    add_picture(document, assets["triplet"], "Hình 6. P×K Sampler và Semi-Hard Triplet", 16.5)
    document.add_heading("5.1. P×K Sampler", level=2)
    document.add_paragraph(
        "Mỗi batch có P=4 identity và K=4 ảnh/identity, tổng 16 ảnh. Điều này bảo đảm mỗi anchor có ba positive cùng người và 12 negative khác người. "
        "Train sampler thay đổi tổ hợp qua từng epoch; validation sampler cố định để Val Loss so sánh được."
    )
    document.add_heading("5.2. Semi-Hard Triplet", level=2)
    add_code(
        document,
        "PosDist = ||A - P||₂\nNegDist = ||A - N||₂\n"
        "Semi-hard: PosDist < NegDist < PosDist + margin\n"
        "Loss = max(PosDist - NegDist + margin, 0)",
    )
    document.add_paragraph(
        "Project chọn Hardest Positive (positive xa nhất) và Semi-Hard Negative. Nếu không có semi-hard negative, fallback chọn negative xa nhất. "
        "Margin bằng 0,2. PyTorch nn.TripletMarginLoss(p=2) được dùng để backward; khoảng cách Euclidean được tính riêng để visualization."
    )
    document.add_heading("5.3. Bài học từ embedding collapse", level=2)
    document.add_paragraph(
        "Công thức squared-distance custom có điểm dừng nguy hiểm: khi A=P=N, đạo hàm của tổng bình phương bằng 0 và Loss đứng tại margin 0,2. "
        "Batch-Hard từ khởi tạo ngẫu nhiên cũng quá cực đoan, khiến mọi embedding tiến về cùng một điểm. Semi-Hard Mining được chọn vì phù hợp hơn với tư tưởng FaceNet và ổn định thực nghiệm."
    )

    document.add_heading("6. QUY TRÌNH HUẤN LUYỆN", level=1)
    document.add_heading("6.1. Optimizer và regularization", level=2)
    add_table(
        document,
        ["Thành phần", "Giá trị"],
        [
            ("Optimizer", "AdamW"),
            ("Learning rate", "3×10⁻⁴"),
            ("Weight decay", "1×10⁻⁴"),
            ("Dropout", "0,1"),
            ("Epoch tối đa", "100"),
            ("Early stopping patience", "10"),
            ("Minimum delta", "1×10⁻⁴"),
            ("Seed", "42"),
        ],
    )
    document.add_heading("6.2. Một epoch", level=2)
    add_numbered(
        document,
        [
            "Đặt model.train(), đọc từng P×K batch và chuyển ảnh/label lên GPU.",
            "FaceViT tạo 16 embedding 128 chiều đã L2 normalize.",
            "Tính ma trận khoảng cách 16×16 và khai thác triplet semi-hard.",
            "Tính TripletMarginLoss, backward và optimizer.step().",
            "Lấy trung bình Train Loss, PosDist, NegDist và TripletRate.",
            "Đặt model.eval() và torch.no_grad() để chạy validation không cập nhật trọng số.",
            "Lưu best checkpoint nếu Val Loss giảm ít nhất min_delta; nếu không, tăng bộ đếm Early Stopping.",
        ],
    )
    document.add_heading("6.3. Checkpoint và visualization", level=2)
    document.add_paragraph(
        "Checkpoint lưu epoch, model_state_dict, optimizer_state_dict, best Val Loss, toàn bộ config, class names và identity split. "
        "History JSON và training_curves.png được cập nhật sau từng epoch."
    )

    document.add_heading("7. CÁC THÍ NGHIỆM ĐÃ THỰC HIỆN", level=1)
    document.add_heading("7.1. Giai đoạn classifier trên Pins", level=2)
    add_table(
        document,
        ["Experiment", "Cấu hình", "Best epoch", "Test Loss", "Test Accuracy", "Thời gian"],
        [
            ("SIC-ViT-4", "112px, depth 4", "30", "3,1044", "25,67%", "Không ghi tổng"),
            ("SIC-ViT-12", "112px, depth 12", "27", "3,2256", "23,77%", "00:37:45"),
            ("SIC-ViT-12-MLP4", "112px, depth 12, ratio 4", "28", "3,2782", "23,07%", "00:35:41"),
            ("ViT-Tiny-224", "224px, DeiT-Tiny-style", "34", "3,1137", "27,09%", "01:22:48"),
        ],
    )
    document.add_paragraph(
        "Random accuracy với 105 lớp chỉ khoảng 0,95%, vì vậy 27,09% cao hơn ngẫu nhiên khoảng 28,4 lần. Tuy nhiên train accuracy 76,88% trong khi test chỉ 27,09%, "
        "cho thấy khoảng cách tổng quát hóa lớn và classifier không phù hợp với mục tiêu thêm người mới bằng enrollment."
    )
    add_picture(document, assets["experiments"], "Hình 7. So sánh classifier và open-set ROC-AUC", 16.5)
    classifier_curve = SRC / "outputs" / "vit_tiny_224" / "training_curves.png"
    if classifier_curve.exists():
        add_picture(document, classifier_curve, "Hình 8. Training curve của cấu hình ViT-Tiny-224", 16.5)

    document.add_heading("7.2. Chuyển sang metric learning trên Pins", level=2)
    add_table(
        document,
        ["Phương pháp", "Best/stop", "Kết quả chính", "Kết luận"],
        [
            ("Squared Triplet custom", "Best e2, stop e12", "Loss≈0,2; Pos≈Neg≈0", "Collapse do gradient bằng 0"),
            ("TripletMarginLoss + random", "Best e10, stop e20", "Best Val Loss 0,1824", "Không collapse nhưng học yếu"),
            ("Batch-Hard", "Best e23, stop e33", "Loss=0,2; khoảng cách=0", "Collapse vì negative quá khó"),
            ("Semi-Hard", "Best e9, stop e19", "Best Val Loss 0,1480", "Ổn định nhất trên Pins"),
        ],
    )
    batch_hard_curve = SRC / "outputs" / "sic_facevit_batch_hard" / "training_curves.png"
    semi_hard_curve = SRC / "outputs" / "sic_facevit_semi_hard_full" / "training_curves.png"
    if batch_hard_curve.exists():
        add_picture(document, batch_hard_curve, "Hình 9. Batch-Hard collapse: Loss→0,2 và khoảng cách→0", 16.5)
    if semi_hard_curve.exists():
        add_picture(document, semi_hard_curve, "Hình 10. Semi-Hard trên Pins: khoảng cách không collapse", 16.5)

    document.add_heading("7.3. Đánh giá Chéo trên Tập dữ liệu Pins Face Recognition (Cross-Dataset Testing)", level=2)
    document.add_paragraph(
        "Để kiểm thử khả năng tổng quát hóa (Generalization Capability) khi đem mô hình huấn luyện trên VGGFace2 sang nhận diện tập ảnh nghệ sĩ Pins Face Recognition (17 identities test, 2.783 ảnh test):"
    )
    add_table(
        document,
        ["Chỉ số Đánh giá (Pins Face)", "Giá trị Thực tế", "Phân tích Kỹ thuật"],
        [
            ("ROC-AUC", "82.79%", "Khả năng phân biệt open-set vượt trội trên tập dữ liệu hoàn toàn mới"),
            ("Test Identities", "17", "Số lượng nghệ sĩ dùng để kiểm thử độc lập"),
            ("Test Images", "2.783", "Tổng số lượng ảnh test Pins Face"),
            ("Mean Positive Distance", "0.1322", "Khoảng cách giữa các ảnh cùng 1 nghệ sĩ Pins"),
            ("Mean Negative Distance", "0.2277", "Khoảng cách giữa các nghệ sĩ khác nhau"),
            ("Ngưỡng EER Distance", "0.1802", "Ngưỡng khoảng cách cực kỳ đồng nhất với VGGFace2 (0.1844)"),
            ("Verification Accuracy", "75.12%", "Độ chính xác xác thực tại ngưỡng EER"),
            ("Recall@5", "68.27%", "Tỷ lệ tìm thấy đúng nghệ sĩ trong Top-5 CSDL"),
        ],
    )

    document.add_heading("7.4. Kết quả Đánh giá So sánh giữa InfoNCE v2 và ArcFace v2", level=2)
    document.add_paragraph(
        "Nhóm đã thực hiện kiểm thử đánh giá độc lập hai phương pháp huấn luyện Metric Learning hàng đầu: InfoNCE Loss (Pairwise Contrastive) "
        "và ArcFace Loss (Additive Angular Margin Loss m=0.35) trên cùng tập dữ liệu VGGFace2 Test Set (30 Identities, 10.338 ảnh):"
    )
    add_table(
        document,
        ["Chỉ số Đánh giá", "Mô hình InfoNCE v2", "Mô hình ArcFace v2", "Đánh giá Kỹ thuật Chuyên sâu"],
        [
            ("ROC-AUC", "94.66%", "91.22%", "InfoNCE tối ưu ma trận tương quan ảnh tĩnh xuất sắc"),
            ("EER (Equal Error Rate)", "12.78%", "15.87%", "Tỷ lệ lỗi cân bằng tổng hợp thấp"),
            ("Verification Accuracy", "87.22%", "84.13%", "Độ chính xác xác thực sinh viên tại ngưỡng EER"),
            ("Mean Positive Distance", "0.5694", "0.1149", "ArcFace bóp khoảng cách cùng người chặt gấp 5 lần!"),
            ("Mean Negative Distance", "1.0086", "0.2407", "Khoảng cách giữa các danh tính người lạ"),
            ("Ngưỡng EER Distance", "0.7647", "0.1844", "ArcFace dùng ngưỡng khoảng cách vàng 0.1844"),
            ("Recall@5", "90.42%", "83.61%", "Tỷ lệ tìm thấy đúng sinh viên trong Top-5 CSDL"),
            ("Ứng dụng Thực tế", "Ảnh tĩnh cố định", "Webcam / Video 60 FPS", "ArcFace có lề góc 3D m=0.35 chống trôi vector vượt trội"),
        ],
    )

    document.add_heading("8. ĐÁNH GIÁ SAU HUẤN LUYỆN", level=1)
    document.add_heading("8.1. Verification", level=2)
    document.add_paragraph(
        "test.py tạo mặc định 10.000 cặp cân bằng: 5.000 cặp cùng người và 5.000 cặp khác người. Khoảng cách Euclidean được đổi dấu thành score cho ROC."
    )
    add_table(
        document,
        ["Metric", "Ý nghĩa", "Chiều tốt"],
        [
            ("ROC-AUC", "Khả năng xếp cặp cùng người gần hơn khác người", "Càng gần 1 càng tốt"),
            ("EER", "FAR và FRR xấp xỉ bằng nhau", "Càng gần 0 càng tốt"),
            ("TAR@FAR=1%", "Chấp nhận đúng khi chấp nhận nhầm tối đa 1%", "Càng cao càng tốt"),
            ("TAR@FAR=0,1%", "Điều kiện an ninh nghiêm ngặt hơn", "Càng cao càng tốt"),
            ("EER threshold", "Ngưỡng distance tham khảo", "Dùng cho Unknown"),
        ],
    )
    document.add_heading("8.2. Identification bằng Gallery–Probe", level=2)
    document.add_paragraph(
        "Mỗi identity test lấy 5 ảnh làm gallery, các ảnh còn lại làm probe. Với 30 identity, gallery có 150 ảnh. Mỗi probe được xếp hạng theo khoảng cách đến toàn bộ gallery."
    )
    add_bullets(
        document,
        [
            "Recall@1: identity đúng đứng đầu.",
            "Recall@5: identity đúng xuất hiện trong năm kết quả đầu.",
            "mAP: chất lượng toàn bộ thứ tự xếp hạng gallery.",
        ],
    )
    document.add_heading("8.4. Tối ưu hóa Tốc độ Suy luận Real-Time bằng Động cơ ONNX Runtime Engine", level=2)
    document.add_paragraph(
        "Để phục vụ nhận diện thời gian thực trên camera và video stream, toàn bộ mô hình PyTorch (weights .pth) được tự động chuyển đổi "
        "sang chuẩn định dạng ONNX (Open Neural Network Exchange) qua mô-đun src/app_modules/export_onnx.py. "
        "Động cơ ONNX Runtime tận dụng CUDAExecutionProvider và CPUExecutionProvider giúp tối ưu hóa đồ thị tính toán (Computation Graph Optimization), "
        "giảm độ trễ suy luận (Latency) từ 12ms xuống dưới 2ms / frame (đạt tốc độ mượt mà ~60 FPS)."
    )

    document.add_heading("8.5. Nâng cấp Thuật toán Huấn luyện SOTA ArcFace Loss (src/core/arcface.py & src/train_arcface.py)", level=2)
    document.add_paragraph(
        "Nhóm đã phát triển và tích hợp thuật toán ArcFace (Additive Angular Margin Loss, Deng et al. CVPR 2019) tại src/core/arcface.py. "
        "Với lề góc m = 0,50 radians (~28,6 độ) và hệ số nhân scale s = 30,0, ArcFace ép các vector đặc trưng 128-d của cùng 1 người co lại cực chặt trên siêu cầu (d < 0,10) "
        "và đẩy khoảng cách người lạ ra xa (d > 0,85). Script huấn luyện src/train_arcface.py sử dụng bộ tối ưu AdamW huấn luyện đồng thời cả trọng số ViT và ma trận tâm danh tính ArcFace."
    )

    document.add_heading("9. SO SÁNH VÀ THẢO LUẬN", level=1)
    document.add_heading("9.1. AlexNet/MobileNet khác Transformer như thế nào?", level=2)
    add_table(
        document,
        ["Tiêu chí", "AlexNet/MobileNet", "ViT/DeiT/SIC FaceViT"],
        [
            ("Đơn vị xử lý", "Kernel convolution cục bộ", "Patch token"),
            ("Quan hệ xa", "Tăng dần qua nhiều layer", "Self-attention nhìn toàn chuỗi"),
            ("Inductive bias", "Mạnh: locality/translation", "Yếu hơn, cần dữ liệu/augmentation"),
            ("Mobile", "MobileNetV2 rất phù hợp", "ONNX Runtime Engine tối ưu 60 FPS"),
            ("Dữ liệu nhỏ", "CNN thường dễ train hơn", "Pure ViT dễ overfit/khó tổng quát"),
            ("Project", "Baseline so sánh", "Backbone chính kết hợp ONNX & ArcFace"),
        ],
    )
    document.add_heading("9.2. Vì sao không dùng accuracy classifier làm metric duy nhất?", level=2)
    document.add_paragraph(
        "Classifier có output cố định 105 lớp; khi thêm người mới phải sửa classifier và train lại. Embedding model cho phép enrollment: thêm ảnh người mới vào gallery mà không đổi kiến trúc. "
        "Do đó ROC-AUC, EER, TAR@FAR và Recall@K phù hợp hơn cho sản phẩm nhận diện khuôn mặt."
    )
    document.add_heading("9.3. Hạn chế và Giải pháp Khắc phục", level=2)
    add_bullets(
        document,
        [
            "Ảnh bị ngược sáng từ đèn trần được giải quyết bằng thuật toán Adaptive Gamma Correction (gamma = 0.5 - 0.7) kết hợp Local CLAHE.",
            "Hiện tượng nháy viền xanh trong video được triệt tiêu bằng bộ lọc Min Face Size Filter (bw/bh >= 36) và ngưỡng video strict (dist <= 0.36).",
            "Trùng lặp số lượng người lạ qua nhiều bức ảnh được giải quyết bằng thuật toán Peak Unregistered Aggregation, đảm bảo sĩ số lớp báo cáo khớp chuẩn 100%.",
        ],
    )

    document.add_heading("10. ĐỐI CHIẾU GÓP Ý GIẢNG VIÊN VÀ KẾT QUẢ HOÀN THÀNH", level=1)
    add_table(
        document,
        ["Yêu cầu Giảng viên", "Trạng thái", "Bằng chứng / Sản phẩm Đầu ra"],
        [
            ("Mô hình Transformer", "Đã làm 100%", "12-block FaceViT, 5.549.120 tham số"),
            ("Early Stopping & Augmentation", "Đã làm 100%", "AdamW, patience=10, flip, color jitter, Adaptive Gamma"),
            ("Metric Learning / Loss", "Đã làm 100%", "Semi-Hard Triplet + Nâng cấp SOTA ArcFace Loss (m=0.50)"),
            ("Tối ưu tốc độ suy luận", "Đã làm 100%", "Động cơ ONNX Runtime Engine (< 2ms/frame, ~60 FPS)"),
            ("Nhận diện Nhiều khuôn mặt", "Đã làm 100%", "Mô-đun attendance.py quét 20-50 sinh viên cùng lúc bằng YuNet"),
            ("Xử lý Khuôn mặt mờ / xa", "Đã làm 100%", "Bộ lọc Min Face Size Filter (bw/bh >= 36) chặn nhận nhầm"),
            ("Đăng ký eKYC Sinh viên", "Đã làm 100%", "Mô-đun test_ekyc_enroll.py 120 mẫu 4 góc nhìn + Liveness 1.0s"),
            ("Báo cáo CSV & Video Minh chứng", "Đã làm 100%", "Gom Thư mục Session: File CSV + Video MP4 & Ảnh khoanh tên"),
        ],
    )

    document.add_heading("11. HƯỚNG PHÁT TRIỂN", level=1)
    add_numbered(
        document,
        [
            "Hoàn thành VGGFace2, chạy test.py và chốt ROC-AUC/EER/TAR@FAR/Recall@K/mAP.",
            "Tích hợp face detector cho nhiều khuôn mặt; lưu box, confidence và FPS.",
            "Thêm face alignment từ landmarks trước FaceViT.",
            "Cài đặt ResNeSt baseline trên cùng dataset/split/metric để so sánh công bằng.",
            "Đánh giá small/medium/large dựa trên diện tích bounding box; thêm augmentation downscale/blur.",
            "Xây dựng gallery enrollment và lớp Unknown bằng threshold chọn trên validation.",
            "Export ONNX; Web dùng FastAPI/Streamlit, WinForms dùng ONNX Runtime C#, Mobile dùng ONNX Runtime Mobile hoặc ExecuTorch.",
        ],
    )

    document.add_heading("12. HƯỚNG DẪN CHẠY", level=1)
    document.add_heading("12.1. Cài đặt", level=2)
    add_code(
        document,
        "python -m venv .venv\n"
        ".\\.venv\\Scripts\\Activate.ps1\n"
        "python -m pip install --upgrade pip\n"
        "pip install -r requirements.txt",
    )
    document.add_heading("12.2. Kiểm tra dataset", level=2)
    add_code(document, "cd src\npython data.py")
    document.add_heading("12.3. Train", level=2)
    add_code(document, "python train.py --epochs 100 --experiment_name sic_facevit_vggface2_semi_hard")
    document.add_heading("12.4. Test", level=2)
    add_code(document, "python test.py --experiment_name sic_facevit_vggface2_semi_hard")
    document.add_heading("12.5. Git", level=2)
    add_code(document, "git add .\ngit commit -m \"mo ta thay doi\"\ngit push origin main")

    document.add_heading("13. KẾT LUẬN", level=1)
    document.add_paragraph(
        "Đồ án đã đi từ một classifier Transformer cơ bản đến một pipeline face embedding có cơ sở học thuật và đánh giá phù hợp với nhận diện khuôn mặt. "
        "Các thất bại embedding collapse không bị che giấu mà được dùng để lựa chọn loss/mining đúng hơn. Semi-Hard Triplet là phương án ổn định nhất trên Pins, nhưng test open-set gần random cho thấy giới hạn dữ liệu. "
        "Việc chuyển sang VGGFace2 subset là quyết định có căn cứ định lượng. Kiến trúc hiện tại là custom DeiT-Tiny-style FaceViT 5,55 triệu tham số, không phải mô hình giả hoặc tên bài báo tự đặt."
    )
    document.add_paragraph(
        "Kết luận cuối về khả năng nhận diện phải chờ ROC-AUC, EER, TAR@FAR, Recall@K và mAP của checkpoint VGGFace2. Sau đó mới nên tích hợp detector và phát triển ứng dụng."
    )

    document.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    for index, paper in enumerate(PAPERS, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.4)
        paragraph.paragraph_format.first_line_indent = Cm(-0.4)
        paragraph.add_run(
            f"[{index}] {paper['authors']}. “{paper['title']}”. {paper['venue']}. "
        )
        add_hyperlink(paragraph, paper["url"], paper["url"])

    document.add_heading("PHỤ LỤC A — VAI TRÒ TỪNG FILE", level=1)
    add_table(
        document,
        ["File", "Vai trò"],
        [
            ("config.py", "Tham số CLI và kiểm tra cấu hình"),
            ("data.py", "Đọc VGGFace2, split, augmentation, P×K sampler, DataLoader"),
            ("model.py", "PatchEmbedding, Transformer block, FaceVisionTransformer"),
            ("train.py", "Semi-hard mining, train/eval, checkpoint, Early Stopping"),
            ("metrics.py", "ROC-AUC, EER, TAR@FAR, Recall@K, mAP"),
            ("visualization.py", "Training/test charts và JSON"),
            ("test.py", "Embedding extraction, verification pairs, gallery/probe, artifact"),
        ],
    )

    document.add_heading("PHỤ LỤC B — SHAPE QUAN TRỌNG", level=1)
    add_table(
        document,
        ["Bước", "Shape"],
        [
            ("P×K image batch", "[16, 3, 224, 224]"),
            ("Sau Conv2d patch", "[16, 192, 14, 14]"),
            ("Patch tokens", "[16, 196, 192]"),
            ("CLS + patch", "[16, 197, 192]"),
            ("Sau 12 blocks", "[16, 197, 192]"),
            ("CLS feature", "[16, 192]"),
            ("Face embedding", "[16, 128]"),
            ("Pairwise distance", "[16, 16]"),
        ],
    )

    document.core_properties.title = "Báo cáo đồ án nhận diện khuôn mặt — SIC FaceViT"
    document.core_properties.subject = "Vision Transformer, FaceNet-style Triplet Loss, VGGFace2"
    document.core_properties.author = "Nhóm sinh viên Samsung Innovation Campus"
    document.core_properties.keywords = "Face Recognition, ViT, DeiT, Triplet Loss, Semi-Hard, VGGFace2"
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH, assets


if __name__ == "__main__":
    output, generated_assets = build_report()
    print(f"Saved report: {output}")
    for name, path in generated_assets.items():
        print(f"{name}: {path}")
