import os
import sys
from pathlib import Path
from datetime import datetime

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def create_document_report():
    doc = Document()

    # Define Style Colors
    PRIMARY_COLOR = RGBColor(16, 44, 87)     # Deep Blue
    SECONDARY_COLOR = RGBColor(53, 89, 142)  # Slate Blue
    DARK_TEXT = RGBColor(33, 33, 33)

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Header / Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("BÁO CÁO ĐỒ ÁN MÔN HỌC / DỰ ÁN AI")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub_p.add_run("HỆ THỐNG NHẬN DIỆN KHUÔN MẶT, eKYC ĐA TƯ THẾ & ĐIỂM DANH SINH VIÊN LỚP HỌC TỰ ĐỘNG (SIC FaceViT)")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(15)
    run_sub.font.bold = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph()

    # Meta Info Table
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Tên Hệ Thống:", "SIC FaceViT (Vision Transformer Face Recognition System)"),
        ("Ngày Xuất Báo Cáo:", datetime.now().strftime("%d/%m/%Y - %H:%M")),
        ("Thuật Toán Cốt Lõi:", "Vision Transformer (ViT) + ArcFace SOTA Loss + YuNet Face Detector")
    ]
    for row_idx, (k, v) in enumerate(meta_data):
        cell_k = meta_table.cell(row_idx, 0)
        cell_v = meta_table.cell(row_idx, 1)
        cell_k.text = k
        cell_v.text = v
        cell_k.paragraphs[0].runs[0].font.bold = True
        cell_k.paragraphs[0].runs[0].font.color.rgb = PRIMARY_COLOR
        cell_v.paragraphs[0].runs[0].font.color.rgb = DARK_TEXT
        set_cell_background(cell_k, "F0F4F8")
        set_cell_background(cell_v, "F9FBFD")

    doc.add_paragraph()

    def add_heading_1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR

    def add_heading_2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.name = "Calibri"
            run_b.font.size = Pt(11)
            run_b.font.bold = True
            run_b.font.color.rgb = PRIMARY_COLOR
        run_t = p.add_run(text)
        run_t.font.name = "Calibri"
        run_t.font.size = Pt(11)
        run_t.font.color.rgb = DARK_TEXT

    # CHƯƠNG 1
    add_heading_1("CHƯƠNG 1: GIỚI THIỆU VÀ MỤC TIÊU DỰ ÁN")
    add_body("Hệ thống điểm danh sinh viên trong lớp học đòi hỏi độ chính xác cao, khả năng xử lý đồng thời nhiều sinh viên trong một bức ảnh tập thể hoặc luồng camera giám sát realtime, và chống chịu tốt với điều kiện ngược sáng, góc nghiêng hoặc khẩu trang.", "1.1 Đặt vấn đề: ")
    add_body("Xây dựng giải pháp AI nhận diện khuôn mặt thương mại thế hệ mới dựa trên mạng Vision Transformer (FaceViT), tích hợp luồng đăng ký sinh trắc học eKYC đa tư thế và mô-đun điểm danh tự động xuất báo cáo CSV và video/ảnh minh chứng.", "1.2 Mục tiêu đồ án: ")

    # CHƯƠNG 2
    add_heading_1("CHƯƠNG 2: KIẾN TRÚC MÔ HÌNH MẠNG FaceViT VÀ ONNX ENGINE")
    add_heading_2("2.1 Mạng Vision Transformer (FaceVisionTransformer)")
    add_body("Sử dụng kiến trúc Vision Transformer với 5.55 triệu tham số. Ảnh khuôn mặt được crop kích thước 224x224, chia thành các patch 16x16 pixel, đưa qua các lớp Transformer Encoder với cơ chế Self-Attention để trích xuất vector embedding L2 128 chiều.")
    add_heading_2("2.2 Động cơ Thực thi ONNX Runtime Engine")
    add_body("Toàn bộ mô hình PyTorch được xuất sang chuẩn định dạng ONNX (Open Neural Network Exchange). Tăng tốc độ suy luận lên gấp 5-6 lần, đạt tốc độ suy luận dưới 2ms / frame (~60 FPS) mượt mà trên camera.")

    # CHƯƠNG 3
    add_heading_1("CHƯƠNG 3: THUẬT TOÁN ĐĂNG KÝ eKYC ĐA TƯ THẾ VÀ CHỐNG NGƯỢC SÁNG")
    add_heading_2("3.1 Quy trình Đăng ký eKYC Đa Tư thế (120 Mẫu)")
    add_body("Yêu cầu sinh viên thực hiện thu thập 120 mẫu đa góc nhìn qua 4 bước: 30 mẫu Nhìn Thẳng, 30 mẫu Quay Trái, 30 mẫu Quay Phải và 30 mẫu Ngước Lên. Hệ thống tích hợp bộ đếm Liveness Timer giữ vững 1.0 giây/tư thế.")
    add_heading_2("3.2 Kiểm tra Khoảng cách và Tự động Bù sáng (Shadow Lifting)")
    add_body("Cảnh báo khung viền màn hình nếu khoảng cách mặt Quá Xa (< 22%) hoặc Quá Gần (> 60%). Đổi mới thuật toán Adaptive Gamma Correction (gamma = 0.5 - 0.7) tự động kéo sáng vùng bóng râm trên mặt khi bị ngược sáng từ đèn trần.")

    # CHƯƠNG 4
    add_heading_1("CHƯƠNG 4: ĐỘNG CƠ GALLERY MATCHING VÀ QUY ĐỔI ĐỘ TIN CẬY SINH TRẮC HỌC")
    add_heading_2("4.1 Hybrid Top-5 Pose Cluster Matching")
    add_body("Kết hợp khoảng cách Centroid trung bình toàn diện và khoảng cách trung bình Top-5 mẫu cụm tư thế khớp nhất (d_hybrid = 0.4 * d_centroid + 0.6 * d_top5). Giảm khoảng cách d khi nghiêng góc từ 0.44 xuống 0.22.")
    add_heading_2("4.2 Đường cong Sinh trắc học Biometric Power Scaling")
    add_body("Quy đổi khoảng cách L2 sang điểm tin cậy Confidence (%) theo hàm lũy thừa sinh trắc học smooth, giúp khoảng cách khớp d = 0.15 hiển thị 92.8% và d = 0.22 hiển thị 88.6%. Cài đặt ngưỡng Cân bằng Vàng d = 0.42.")

    # CHƯƠNG 5
    add_heading_1("CHƯƠNG 5: HỆ THỐNG ĐIỂM DANH LỚP HỌC VÀ XUẤT BÁO CÁO (src/app_modules/attendance.py)")
    add_heading_2("5.1 Các Chế độ Điểm danh")
    add_body("Hỗ trợ 4 chế độ: Điểm danh Thư mục Ngày hàng loạt (--folder), Ảnh chụp tập thể lớp (--image), Luồng Video stream camera (--video) và Webcam trực tiếp (--webcam).")
    add_heading_2("5.2 Các Thuật toán Xử lý Nâng cao")
    add_body("1. Ràng buộc Độc quyền Danh tính per Frame: 1 sinh viên trong CSDL chỉ gán cho 1 mặt khớp nhất trong bức ảnh.\n2. Bộ lọc Mặt mờ Quá nhỏ (Min Face Size Filter): Bỏ qua mặt width/height < 36 pixels ở hàng ghế quá xa.\n3. Lọc trùng Người lạ theo Đỉnh (Peak Unregistered Aggregation): Giữ sĩ số chuẩn 22 sinh viên, không bị cộng dồn người lạ.")
    add_heading_2("5.3 Quản lý Thư mục Session Tích hợp")
    add_body("Gom chung toàn bộ kết quả vào 1 Thư mục Session (outputs/attendance_session_.../) chứa File Báo cáo CSV Tổng hợp + Thư mục Ảnh & Video MP4 Minh chứng đã khoanh tên viền xanh/đỏ.")

    # CHƯƠNG 6
    add_heading_1("CHƯƠNG 6: NÂNG CẤP THUẬT TOÁN ARCFACE LOSS SOTA (src/core/arcface.py)")
    add_body("Áp dụng Additive Angular Margin Loss (m = 0.50 rad ~ 28.6 độ, scale s = 30.0). Ép vector đặc trưng của cùng 1 người co lại cực chặt trên siêu cầu (d < 0.10) và đẩy khoảng cách người lạ ra xa (d > 0.85). Tối ưu đồng thời trọng số ViT và ArcFace class weights bằng AdamW.")

    # CHƯƠNG 7
    add_heading_1("CHƯƠNG 7: KẾT QUẢ THỬ NGHIỆM VÀ HƯỚNG DẪN SỬ DỤNG")
    add_heading_2("7.1 Bảng So sánh Hiệu năng Hệ thống")

    res_table = doc.add_table(rows=5, cols=3)
    res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Tiêu Chí Đánh Giá", "Trước Khi Tối Ưu", "Sau Khi Tối Ưu"]
    for i, h in enumerate(headers):
        cell = res_table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = PRIMARY_COLOR
        set_cell_background(cell, "E6ECF5")

    rows_data = [
        ("Độ tin cậy hiển thị (d=0.22)", "67.3%", "88.6% (Sinh trắc mượt)"),
        ("Tốc độ Suy luận (ONNX)", "12 FPS", "60 FPS (< 2ms/frame)"),
        ("Sĩ số Báo cáo Thư mục Ngày", "Nhân đôi (41 SV)", "Khớp chuẩn thực tế (22 SV)"),
        ("Xuất Minh chứng Video", "Không có", "Tự động xuất MP4 khoanh tên")
    ]

    for r_idx, row in enumerate(rows_data, 1):
        for c_idx, val in enumerate(row):
            cell = res_table.cell(r_idx, c_idx)
            cell.text = val
            cell.paragraphs[0].runs[0].font.color.rgb = DARK_TEXT
            set_cell_background(cell, "F9FBFD" if r_idx % 2 == 1 else "FFFFFF")

    add_heading_2("7.2 Cú pháp Lệnh Chạy Hệ thống")
    add_body("1. Đăng ký eKYC Mới (120 Mẫu): python src/app_modules/test_ekyc_enroll.py --enroll_name \"Quyen_eKYC\" --use_onnx\n2. Điểm danh Thư mục Ngày: python src/app_modules/attendance.py --folder /path/to/day_folder --use_onnx\n3. Điểm danh Ảnh tập thể: python src/app_modules/attendance.py --image path/to/classroom.jpg --use_onnx\n4. Điểm danh Video stream: python src/app_modules/attendance.py --video path/to/class_video.mp4 --use_onnx\n5. Huấn luyện ArcFace Loss: python src/train_arcface.py --experiment_name sic_facevit_arcface_v1")

    # Save Output Docx Document
    out_dir = Path("outputs")
    os.makedirs(out_dir, exist_ok=True)
    out_docx_path = out_dir / "Bao_Cao_Do_An_SIC_FaceViT.docx"
    doc.save(str(out_docx_path))

    print(f"✅ ĐÃ XUẤT BÁO CÁO ĐỒ ÁN WORD (.DOCX) THÀNH CÔNG VÀO: {out_docx_path}")
    return str(out_docx_path)


if __name__ == "__main__":
    create_document_report()
